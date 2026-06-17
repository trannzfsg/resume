from docx import Document

path = r"C:\github\resume\tranzha.docx"
doc = Document(path)
texts = [p.text for p in doc.paragraphs if p.text.strip()]
search_text = "\n".join(texts).lower()
required = [
    "Tran Zha",
    "tranzha83@gmail.com",
    "Software Engineering Manager",
    "Full Australian working rights",
    "Australian permanent resident",
    "Key Skills",
    "Harness engineering",
    "reusable skills",
    "Loop engineering",
    "working myself out of the job",
    "ISO 27001",
    "SOC 2 Type II",
    "PII data handling",
    "Loan Market Group",
    "event-sourcing patterns",
    "backward-compatible projections",
    "scope and debt tradeoffs",
    "Drove domain modelling",
    "roadmap alignment",
    "~10 to 30+ engineers",
    "0 to 30+ engineers",
    "Scrum master",
    "Managed cloud spend",
    "incident response",
    "Jul 2019 to Apr 2021",
    "Languages: English, Mandarin",
]
for needle in required:
    if needle.lower() not in search_text:
        raise SystemExit(f"missing: {needle}")
print(f"paragraphs={len(doc.paragraphs)}")
print(f"first={texts[0]}")
print("smoke=ok")
