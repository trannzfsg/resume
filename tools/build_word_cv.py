import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "tranzha.md"
OUT = ROOT / "tranzha.docx"


BLUE = RGBColor(46, 116, 181)
CHARCOAL = RGBColor(34, 34, 34)
MUTED = RGBColor(85, 85, 85)
LINK_BLUE = RGBColor(5, 99, 193)

LINK_RE = re.compile(r"(https?://[^\s|]+|[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,})", re.I)


def set_spacing(paragraph, before=0, after=6, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_run(run, size=10.5, bold=False, color=CHARCOAL, underline=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.underline = underline


def add_border(paragraph, color="D9E2F3", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def strip_markdown(text):
    return text.replace("**", "")


def add_text_runs(paragraph, text, size=10.5, bold=False, color=CHARCOAL):
    pos = 0
    for match in LINK_RE.finditer(text):
        if match.start() > pos:
            set_run(paragraph.add_run(text[pos:match.start()]), size=size, bold=bold, color=color)
        link_text = match.group(0)
        set_run(paragraph.add_run(link_text), size=size, bold=bold, color=LINK_BLUE, underline=True)
        pos = match.end()
    if pos < len(text):
        set_run(paragraph.add_run(text[pos:]), size=size, bold=bold, color=color)


def add_markdown_runs(paragraph, text, size=10.5, color=CHARCOAL):
    pos = 0
    for match in re.finditer(r"\*\*([^*]+)\*\*", text):
        if match.start() > pos:
            add_text_runs(paragraph, text[pos:match.start()], size=size, color=color)
        add_text_runs(paragraph, match.group(1), size=size, bold=True, color=color)
        pos = match.end()
    if pos < len(text):
        add_text_runs(paragraph, text[pos:], size=size, color=color)


def add_section(doc, title):
    p = doc.add_paragraph(style="Heading 1")
    set_spacing(p, before=8, after=3, line=1.0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title.upper())
    set_run(run, size=11.2, bold=True, color=BLUE)
    add_border(p)
    return p


def add_body_paragraph(doc, text, after=4):
    p = doc.add_paragraph()
    set_spacing(p, after=after)
    add_markdown_runs(p, text, size=10.2)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.03
    add_markdown_runs(p, text, size=9.8)
    return p


def add_role(doc, title):
    p = doc.add_paragraph(style="Heading 2")
    set_spacing(p, before=5, after=0, line=1.0)
    p.paragraph_format.keep_with_next = True
    add_markdown_runs(p, title, size=10.2)
    for run in p.runs:
        run.font.bold = True
    return p


def add_project(doc, title):
    p = doc.add_paragraph(style="Heading 3")
    set_spacing(p, before=2, after=1, line=1.0)
    p.paragraph_format.keep_with_next = True
    add_markdown_runs(p, title, size=9.8)
    for run in p.runs:
        run.font.bold = True
    return p


def parse_markdown(markdown):
    blocks = []
    paragraph = []
    bullets = []

    def flush_paragraph():
        nonlocal paragraph
        if paragraph:
            blocks.append(("p", paragraph))
            paragraph = []

    def flush_bullets():
        nonlocal bullets
        if bullets:
            blocks.append(("ul", bullets))
            bullets = []

    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            flush_paragraph()
            flush_bullets()
            continue
        if line.startswith("- "):
            flush_paragraph()
            bullets.append(line[2:].strip())
            continue

        flush_bullets()
        if line.startswith("### "):
            flush_paragraph()
            blocks.append(("h3", line[4:].strip()))
        elif line.startswith("## "):
            flush_paragraph()
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("# "):
            flush_paragraph()
            blocks.append(("h1", line[2:].strip()))
        else:
            paragraph.append(line.removesuffix("  ").strip())

    flush_paragraph()
    flush_bullets()
    return blocks


def configure_document(doc):
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.2)
    styles["Normal"].font.color.rgb = CHARCOAL
    styles["Title"].font.name = "Calibri"
    styles["Title"].font.size = Pt(21)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor(31, 77, 120)
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(11.2)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = BLUE
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(10.2)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = CHARCOAL
    styles["Heading 3"].font.name = "Calibri"
    styles["Heading 3"].font.size = Pt(9.8)
    styles["Heading 3"].font.bold = True
    styles["Heading 3"].font.color.rgb = CHARCOAL
    styles["List Bullet"].font.name = "Calibri"
    styles["List Bullet"].font.size = Pt(9.8)


def set_core_properties(doc):
    props = doc.core_properties
    props.author = "Tran Zha"
    props.title = "Tran Zha CV"
    props.subject = "Software Engineering Manager, Platform Modernisation, Technical Leadership"
    props.keywords = (
        "Software Engineering Manager, Engineering Manager, Engineering Lead, "
        "Technical Lead, Platform Modernisation, AWS, Kubernetes, C# .NET"
    )
    props.comments = "Generated from tranzha.md"


def add_header_block(doc, blocks, index):
    title = blocks[index][1]
    name = doc.add_paragraph(style="Title")
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(name, after=1, line=1.0)
    set_run(name.add_run(title), size=21, bold=True, color=RGBColor(31, 77, 120))

    if index + 1 >= len(blocks) or blocks[index + 1][0] != "p":
        return index + 1

    lines = blocks[index + 1][1]
    if not lines:
        return index + 2

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(subtitle, after=1, line=1.0)
    add_markdown_runs(subtitle, lines[0], size=10.2, color=CHARCOAL)
    for run in subtitle.runs:
        run.font.bold = True

    if len(lines) > 1:
        location = doc.add_paragraph()
        location.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(location, after=1, line=1.0)
        add_markdown_runs(location, lines[1], size=9.5, color=MUTED)

    if len(lines) > 2:
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(contact, after=5, line=1.0)
        add_markdown_runs(contact, " ".join(lines[2:]), size=9.2, color=MUTED)

    return index + 2


def build():
    doc = Document()
    configure_document(doc)
    blocks = parse_markdown(SRC.read_text(encoding="utf-8"))

    index = 0
    if blocks and blocks[0][0] == "h1":
        index = add_header_block(doc, blocks, 0)

    while index < len(blocks):
        kind, value = blocks[index]
        if kind == "h1":
            add_role(doc, value)
        elif kind == "h2":
            add_section(doc, value)
        elif kind == "h3":
            add_role(doc, value)
        elif kind == "p":
            text = " ".join(value)
            plain = strip_markdown(text)
            if re.fullmatch(r"[A-Z][a-z]{2} \d{4} to (Present|[A-Z][a-z]{2} \d{4})", plain):
                p = doc.add_paragraph()
                set_spacing(p, before=0, after=4, line=1.0)
                p.paragraph_format.keep_with_next = True
                set_run(p.add_run(plain), size=9.5, bold=True, color=MUTED)
            elif plain.startswith("Selected initiative:"):
                add_project(doc, plain)
            else:
                add_body_paragraph(doc, text, after=6)
        elif kind == "ul":
            for item in value:
                add_bullet(doc, item)
        index += 1

    set_core_properties(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
